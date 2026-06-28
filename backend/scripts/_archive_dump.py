"""Diagnostic evidence dumper for a single data/archive file.

Usage (from repo root):
    cd backend && PYTHONPATH=. .venv/bin/python -m scripts._archive_dump "<filename-in-data/archive>" [--rows N] [--lines N]

Prints ground-truth evidence: detected format, page/sheet structure, raw text head,
pdfplumber tables (if any), reconstructed word-geometry lines, openpyxl/xlrd sheet
rows, and what the CURRENT parser produced (parse_file) so a reviewer can judge the
gap between source truth and parser output. Read-only; touches nothing.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
ARCHIVE = REPO / "data" / "archive"


def _p(*a):
    print(*a, flush=True)


def dump_pdf(path: Path, n_lines: int) -> None:
    import pdfplumber
    with pdfplumber.open(str(path)) as pdf:
        _p(f"  pages: {len(pdf.pages)}")
        for pno, pg in enumerate(pdf.pages[:3]):
            tabs = pg.extract_tables() or []
            _p(f"  --- page {pno+1}: ruled_tables={len(tabs)} ---")
            if tabs:
                for row in tabs[0][:8]:
                    _p(f"    T| {row}")
            txt = (pg.extract_text() or "").splitlines()
            _p(f"    text lines ({len(txt)}), sample:")
            for ln in txt[:n_lines]:
                _p(f"    L| {ln!r}")
        # word geometry on page where data starts
        try:
            from app.parsers.pdf_text import _pdf_lines, _cells_by_gap
            pg = pdf.pages[min(1, len(pdf.pages) - 1)]
            words = pg.extract_words(use_text_flow=False)
            _p("  --- word-geometry cell split (page 2, first lines) ---")
            for line in _pdf_lines(words)[:n_lines]:
                _p(f"    G| {_cells_by_gap(line)}")
        except Exception as e:  # noqa: BLE001
            _p(f"  (word-geometry dump failed: {e})")


def dump_docx(path: Path, n_lines: int) -> None:
    from io import BytesIO
    from docx import Document
    from app.parsers.docx_parser import accept_revisions, _table_to_2d
    d = Document(BytesIO(accept_revisions(str(path))))
    _p(f"  tables: {len(d.tables)}")
    for tno, t in enumerate(d.tables[:2]):
        rows = _table_to_2d(t)
        _p(f"  --- table {tno+1}: {len(rows)} rows x {len(rows[0]) if rows else 0} cols ---")
        for r in rows[:n_lines]:
            _p(f"    R| {r}")


def dump_xlsx(path: Path, n_lines: int) -> None:
    name = path.name.lower()
    if name.endswith(".xls"):
        import pandas as pd
        sheets = pd.read_excel(str(path), sheet_name=None, header=None, dtype=str, engine="xlrd")
        for sn, df in sheets.items():
            _p(f"  --- sheet {sn!r}: {df.shape[0]} rows x {df.shape[1]} cols ---")
            for r in df.values.tolist()[:n_lines]:
                _p(f"    X| {[('' if (v!=v) else v) for v in r]}")
    else:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        for ws in wb.worksheets:
            _p(f"  --- sheet {ws.title!r}: dims={ws.dimensions} ---")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= n_lines:
                    break
                _p(f"    X| {list(row)}")
        wb.close()


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("filename")
    ap.add_argument("--rows", type=int, default=25, help="parser rows to show")
    ap.add_argument("--lines", type=int, default=25, help="raw lines to show")
    args = ap.parse_args()

    path = ARCHIVE / args.filename
    if not path.is_file():
        # allow partial / case-insensitive match
        cands = [p for p in ARCHIVE.iterdir() if args.filename.lower() in p.name.lower()]
        if len(cands) == 1:
            path = cands[0]
        else:
            _p(f"NOT FOUND: {args.filename!r}; candidates: {[p.name for p in cands]}")
            return 2

    from app.parsers import detect_format, parse_file
    fmt = detect_format(str(path), path.name)
    _p("=" * 100)
    _p(f"FILE: {path.name}")
    _p(f"DETECTED FORMAT: {fmt.value}   size={path.stat().st_size} bytes")
    _p("=" * 100)

    _p("\n[RAW SOURCE EVIDENCE]")
    v = fmt.value
    try:
        if v in ("pdf", "scan_pdf"):
            dump_pdf(path, args.lines)
        elif v == "docx":
            dump_docx(path, args.lines)
        elif v in ("xlsx", "xls"):
            dump_xlsx(path, args.lines)
        else:
            _p(f"  (no dumper for format {v})")
    except Exception as e:  # noqa: BLE001
        import traceback
        _p(f"  RAW DUMP ERROR: {e}")
        traceback.print_exc()

    _p("\n[CURRENT PARSER OUTPUT  (parse_file)]")
    doc = parse_file(str(path), path.name)
    priced = sum(1 for r in doc.rows if (r.price_resident or r.price_nonresident or r.price_original))
    _p(f"  total_rows={len(doc.rows)}  priced_rows={priced}  used_ocr={doc.used_ocr}  raw_text_len={len(doc.raw_text)}")
    _p(f"  hints: partner={doc.partner_name_hint!r} city={doc.city_hint!r} date={doc.effective_date_hint} bin={doc.bin_hint}")
    if doc.warnings:
        _p("  warnings:")
        for w in doc.warnings[:12]:
            _p(f"    ! {w}")
    _p(f"  first {args.rows} parsed rows (name | res | nonres | code | section):")
    for r in doc.rows[:args.rows]:
        sec = (r.extra or {}).get("section")
        _p(f"    P| {r.service_name_raw!r:55} | {r.price_resident} | {r.price_nonresident} | "
           f"{r.service_code_source!r} | sec={sec!r}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
