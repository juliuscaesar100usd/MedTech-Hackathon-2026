"""Generate a realistic ZIP archive of Kazakhstan clinic price lists.

Produces ``backend/sample_data/archive.zip`` with ~8-10 price-list files that
exercise every parser path and validation rule in MedArchive:

  * 2-3 TEXT PDFs (reportlab, Cyrillic via DejaVuSans TTF) — header block +
    "Наименование услуги | Цена резидент (₸) | Цена нерезидент (₸)" table.
  * 1 genuine SCANNED PDF (image-only) — a price list rendered to a PIL image
    and embedded full-page via PyMuPDF, so the PDF has NO selectable text. OCR
    is verified against the bundled tessdata before packing.
  * 2 XLSX (openpyxl) — multi-sheet, header NOT on the first row (title/merged
    rows above). One has separate resident/non-resident columns; the other adds
    a USD-priced sheet to exercise currency conversion.
  * 1-2 DOCX (python-docx) — header block + services table. One carries real
    TRACKED CHANGES (<w:ins>/<w:del> in word/document.xml) such that accepting
    revisions yields the NEW prices.

Cross-file features:
  * PRICE HISTORY: one clinic emits two price lists (2024-01 and 2025-06) for
    overlapping services with at least one >50% jump (anomaly trigger).
  * EDGE CASES: "договорная" / empty price cells, off-catalog services
    (unmatched + needs_review queues), and rows where non-resident < resident.

Run: ``python -m scripts.generate_sample_archive`` (from ``backend/``).
"""
from __future__ import annotations

import io
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from lxml import etree
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet

import fitz  # PyMuPDF

from ._sampledata_spec import CLINICS, SERVICES, ClinicSpec, ServiceSpec

# --------------------------------------------------------------------------- #
# Paths / fonts.                                                              #
# --------------------------------------------------------------------------- #
OUT_DIR = Path(__file__).resolve().parent.parent / "sample_data"
ARCHIVE_PATH = OUT_DIR / "archive.zip"

DEJAVU = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
DEJAVU_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

_FONT_REGISTERED = False


def _register_fonts() -> None:
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    pdfmetrics.registerFont(TTFont("DejaVu", DEJAVU))
    pdfmetrics.registerFont(TTFont("DejaVu-Bold", DEJAVU_BOLD))
    _FONT_REGISTERED = True


# --------------------------------------------------------------------------- #
# Helpers.                                                                     #
# --------------------------------------------------------------------------- #
def _by_name(name: str) -> ServiceSpec:
    for s in SERVICES:
        if s.name == name:
            return s
    raise KeyError(name)


def _kzt(value: int) -> str:
    """Format a KZT integer with ASCII-space thousands grouping, e.g. '12 500'.

    The parser only strips ASCII space (0x20), so we MUST use a plain space
    (not NBSP/thin space) for the grouping or prices would mis-parse.
    """
    return f"{value:,}".replace(",", chr(32))


def _safe_filename(clinic: ClinicSpec, suffix: str, d: date) -> str:
    """Build a clinic+date filename, e.g. 'Клиника_Сункар_прайс_2025-01-15.pdf'."""
    base = (
        clinic.name.replace("«", "")
        .replace("»", "")
        .replace("Медицинский центр ", "Клиника_")
        .replace("Диагностический центр ", "ДЦ_")
        .replace("Лаборатория ", "Лаб_")
        .replace("Стоматологическая клиника ", "Стом_")
        .replace("Клиника ", "Клиника_")
        .strip()
        .replace(" ", "_")
    )
    return f"{base}_прайс_{d.isoformat()}.{suffix}"


def _header_lines(clinic: ClinicSpec, d: date) -> list[str]:
    return [
        clinic.name,
        clinic.address,
        f"БИН: {clinic.bin}",
        f"тел.: {clinic.phone}    e-mail: {clinic.email}",
        f"Прайс-лист от {d.strftime('%d.%m.%Y')}",
    ]


# A row for a price table: (name, resident, nonresident) where prices are
# pre-formatted strings (so we can inject "договорная"/"" edge cases).
PriceRow = tuple[str, str, str]


def _rows_for(
    names: list[str],
    *,
    res_mult: float = 1.0,
    nonres_mult: float = 1.45,
    use_synonym: bool = False,
) -> list[PriceRow]:
    """Build standard priced rows from catalog services.

    ``use_synonym`` emits an abbreviation/alias instead of the official name so
    the matcher's synonym/fuzzy paths get exercised too.
    """
    out: list[PriceRow] = []
    for nm in names:
        s = _by_name(nm)
        label = s.synonyms[0] if (use_synonym and s.synonyms) else s.name
        res = int(round(s.base * res_mult))
        nonres = int(round(s.base * nonres_mult))
        out.append((label, _kzt(res), _kzt(nonres)))
    return out


# --------------------------------------------------------------------------- #
# TEXT PDFs (reportlab).                                                       #
# --------------------------------------------------------------------------- #
def build_text_pdf(
    clinic: ClinicSpec,
    d: date,
    rows: list[PriceRow],
    *,
    title_note: str | None = None,
) -> bytes:
    _register_fonts()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
        title=f"{clinic.name} прайс {d.isoformat()}",
    )
    styles = getSampleStyleSheet()
    h = ParagraphStyle(
        "H", parent=styles["Normal"], fontName="DejaVu-Bold",
        fontSize=14, leading=18, spaceAfter=2,
    )
    meta = ParagraphStyle(
        "M", parent=styles["Normal"], fontName="DejaVu",
        fontSize=9.5, leading=13,
    )
    note = ParagraphStyle(
        "N", parent=styles["Normal"], fontName="DejaVu",
        fontSize=9, leading=12, textColor=colors.grey,
    )
    cell = ParagraphStyle(
        "C", parent=styles["Normal"], fontName="DejaVu", fontSize=9, leading=11,
    )
    cell_hdr = ParagraphStyle(
        "CH", parent=styles["Normal"], fontName="DejaVu-Bold",
        fontSize=9, leading=11, textColor=colors.white,
    )

    story: list = []
    lines = _header_lines(clinic, d)
    story.append(Paragraph(lines[0], h))
    for ln in lines[1:]:
        story.append(Paragraph(ln, meta))
    if title_note:
        story.append(Spacer(1, 3))
        story.append(Paragraph(title_note, note))
    story.append(Spacer(1, 8))

    data = [[
        Paragraph("Наименование услуги", cell_hdr),
        Paragraph("Цена резидент (₸)", cell_hdr),
        Paragraph("Цена нерезидент (₸)", cell_hdr),
    ]]
    for nm, res, nonres in rows:
        data.append([
            Paragraph(nm, cell),
            Paragraph(res, cell),
            Paragraph(nonres, cell),
        ])

    table = Table(data, colWidths=[105 * mm, 35 * mm, 35 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f6f8b")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9bbfca")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (1, 0), (-1, 0), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#eef5f7")]),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    doc.build(story)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# SCANNED PDF (image-only, OCR target).                                        #
# --------------------------------------------------------------------------- #
def render_pricelist_image(
    clinic: ClinicSpec, d: date, rows: list[PriceRow]
) -> Image.Image:
    """Render a clean, high-contrast price list onto a white image for OCR."""
    W, H = 1654, 2339  # ~A4 @ 200 DPI
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)

    f_title = ImageFont.truetype(DEJAVU_BOLD, 40)
    f_meta = ImageFont.truetype(DEJAVU, 28)
    f_hdr = ImageFont.truetype(DEJAVU_BOLD, 30)
    f_cell = ImageFont.truetype(DEJAVU, 30)

    margin = 90
    y = 80
    lines = _header_lines(clinic, d)
    draw.text((margin, y), lines[0], fill="black", font=f_title)
    y += 64
    for ln in lines[1:]:
        draw.text((margin, y), ln, fill="black", font=f_meta)
        y += 42
    y += 26

    # Column x-positions: name | resident | non-resident.
    col_name = margin
    col_res = 1000
    col_nonres = 1320

    draw.text((col_name, y), "Наименование услуги", fill="black", font=f_hdr)
    draw.text((col_res, y), "Резидент ₸", fill="black", font=f_hdr)
    draw.text((col_nonres, y), "Нерезидент ₸", fill="black", font=f_hdr)
    y += 46
    draw.line((margin, y, W - margin, y), fill="black", width=2)
    y += 18

    for nm, res, nonres in rows:
        draw.text((col_name, y), nm, fill="black", font=f_cell)
        draw.text((col_res, y), res, fill="black", font=f_cell)
        draw.text((col_nonres, y), nonres, fill="black", font=f_cell)
        y += 50

    return img


def image_to_scanned_pdf(img: Image.Image) -> bytes:
    """Embed a PIL image as a full-page image in a PDF (no selectable text).

    The render is near-binary (black text on white), so a palette PNG with
    optimisation compresses it from ~11 MB to a fraction without harming OCR.
    """
    png = io.BytesIO()
    # Quantize to a small palette: huge size win, OCR stays crisp.
    img.convert("L").save(png, format="PNG", optimize=True)
    png_bytes = png.getvalue()

    pdf = fitz.open()
    page = pdf.new_page(width=595, height=842)  # A4 in points
    page.insert_image(page.rect, stream=png_bytes)
    out = pdf.tobytes()
    pdf.close()
    return out


def _rasterize_pdf_page(pdf_bytes: bytes, page_no: int = 0) -> Image.Image:
    """Rasterize a PDF page to a PIL image at the project's OCR DPI.

    Identical to ``app.parsers.pdf_scan``'s rasterization, so OCR verification
    sees exactly what the parser will see.
    """
    from app.config import settings

    zoom = settings.ocr_dpi / 72.0
    with fitz.open(stream=pdf_bytes, filetype="pdf") as d:
        page = d[page_no]
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        return Image.frombytes("RGB", (pix.width, pix.height), pix.samples)


def verify_scanned_ocr(
    img: Image.Image, expect_tokens: list[str]
) -> tuple[bool, str]:
    """OCR the rendered image with the project's tessdata/langs and assert tokens."""
    import os

    import pytesseract

    from app.config import settings

    tdir = str(settings.tessdata_prefix)
    os.environ["TESSDATA_PREFIX"] = tdir
    config = f'--tessdata-dir "{tdir}"'
    text = pytesseract.image_to_string(img, lang=settings.ocr_langs, config=config)
    low = text.lower()
    ok = all(tok.lower() in low for tok in expect_tokens)
    return ok, text


# --------------------------------------------------------------------------- #
# XLSX (openpyxl) — header not on first row.                                   #
# --------------------------------------------------------------------------- #
def _write_xlsx_sheet(
    ws,
    clinic: ClinicSpec,
    d: date,
    headers: list[str],
    rows: list[list[str]],
    *,
    note: str | None = None,
) -> None:
    """Write a sheet whose header is NOT the first row (2-3 title rows above)."""
    ncol = len(headers)
    bold = Font(bold=True)

    # Title rows (merged across all columns) — header lands on row 3 or 4.
    ws.append([clinic.name])
    ws.merge_cells(
        start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=ncol
    )
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=13)

    ws.append([f"БИН: {clinic.bin}   тел.: {clinic.phone}   {clinic.email}"])
    ws.merge_cells(
        start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=ncol
    )

    ws.append([f"Прайс-лист от {d.strftime('%d.%m.%Y')}"])
    ws.merge_cells(
        start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=ncol
    )

    if note:
        ws.append([note])
        ws.merge_cells(
            start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=ncol
        )

    # Header row.
    ws.append(headers)
    for c in range(1, ncol + 1):
        ws.cell(row=ws.max_row, column=c).font = bold

    for r in rows:
        ws.append(r)

    # Reasonable column widths.
    for c in range(1, ncol + 1):
        ws.column_dimensions[get_column_letter(c)].width = 42 if c == 1 else 20


def build_xlsx_two_price(clinic: ClinicSpec, d: date) -> bytes:
    """Multi-sheet workbook with separate resident/non-resident columns (KZT)."""
    wb = Workbook()
    lab = wb.active
    lab.title = "Лаборатория"
    diag = wb.create_sheet("Диагностика")

    lab_names = [
        "Общий анализ крови", "Общий анализ мочи", "Биохимический анализ крови",
        "Глюкоза крови", "Липидный профиль", "Тиреотропный гормон",
        "Витамин D (25-OH)", "С-реактивный белок", "Ферритин",
        "Группа крови и резус-фактор",
    ]
    diag_names = [
        "УЗИ органов брюшной полости", "УЗИ щитовидной железы", "УЗИ почек",
        "Рентгенография органов грудной клетки", "Гастроскопия",
        "Эхокардиография", "Электрокардиография",
    ]

    headers = ["Наименование услуги", "Цена резидент", "Цена нерезидент"]

    lab_rows = [[nm, r, nr] for nm, r, nr in _rows_for(lab_names)]
    # Edge cases inside the lab sheet:
    # 1) non-resident < resident (deliberate inversion)
    lab_rows.append(["Анализ кала на скрытую кровь", _kzt(2800), _kzt(2000)])
    # 2) "договорная" (no number)
    lab_rows.append(["Антитела к COVID-19 IgG", "договорная", "договорная"])
    # 3) empty price cell
    lab_rows.append(["ПЦР на SARS-CoV-2", "", _kzt(11600)])
    # 4) off-catalog services: one clearly UNMATCHED, one in the needs_review band
    lab_rows.append(["Подбор ортопедических стелек", _kzt(5200), _kzt(7500)])
    lab_rows.append(["Иглорефлексотерапия (сеанс)", _kzt(4000), _kzt(5800)])

    diag_rows = [[nm, r, nr] for nm, r, nr in _rows_for(diag_names, use_synonym=True)]

    _write_xlsx_sheet(lab, clinic, d, headers, lab_rows)
    _write_xlsx_sheet(diag, clinic, d, headers, diag_rows)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_xlsx_usd(clinic: ClinicSpec, d: date) -> bytes:
    """Multi-sheet workbook including a USD-priced sheet (currency conversion)."""
    wb = Workbook()
    kzt_sheet = wb.active
    kzt_sheet.title = "Стоматология (тг)"
    usd_sheet = wb.create_sheet("Импланты (USD)")

    dent_names = [
        "Лечение кариеса", "Удаление зуба",
        "Профессиональная гигиена полости рта", "Лечение пульпита",
        "Снимок зуба (прицельный)",
    ]
    kzt_rows = [[nm, r, nr] for nm, r, nr in _rows_for(dent_names)]
    headers_kzt = ["Наименование услуги", "Цена резидент", "Цена нерезидент"]
    _write_xlsx_sheet(kzt_sheet, clinic, d, headers_kzt, kzt_rows)

    # USD sheet: a price column "Цена, USD" + an explicit "Валюта" column so the
    # parser's currency-column mapping path (not just symbol detection) fires.
    usd_rows = [
        ["Установка коронки", "90", "USD"],
        ["Отбеливание зубов", "80", "USD"],
        # Off-catalog service priced in USD -> unmatched + currency conversion.
        ["Газоразрядная визуализация ауры", "1200", "USD"],
    ]
    headers_usd = ["Наименование услуги", "Цена, USD", "Валюта"]
    _write_xlsx_sheet(
        usd_sheet, clinic, d, headers_usd, usd_rows,
        note="Цены в сегменте имплантации указаны в долларах США (USD).",
    )

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# DOCX (python-docx) + tracked changes.                                        #
# --------------------------------------------------------------------------- #
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx_header_and_table(
    clinic: ClinicSpec, d: date, rows: list[PriceRow]
) -> bytes:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "DejaVu Sans"
    style.font.size = Pt(10)

    document.add_heading(clinic.name, level=1)
    for ln in _header_lines(clinic, d)[1:]:
        document.add_paragraph(ln)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Наименование услуги"
    hdr[1].text = "Цена резидент (₸)"
    hdr[2].text = "Цена нерезидент (₸)"
    for nm, res, nonres in rows:
        cells = table.add_row().cells
        cells[0].text = nm
        cells[1].text = res
        cells[2].text = nonres

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_docx_plain(clinic: ClinicSpec, d: date) -> bytes:
    names = [
        "Консультация терапевта", "Консультация кардиолога",
        "Консультация невролога", "Консультация эндокринолога",
        "Консультация гинеколога", "Электрокардиография", "Эхокардиография",
    ]
    rows = list(_rows_for(names))
    # Edge case: non-resident < resident on one consultation row.
    rows.append(("Консультация педиатра", _kzt(5500), _kzt(4800)))
    return _docx_header_and_table(clinic, d, rows)


def build_docx_tracked(clinic: ClinicSpec, d: date) -> tuple[bytes, dict]:
    """Build a DOCX, then inject real tracked changes into word/document.xml.

    Two price cells are revised: the OLD value is wrapped in <w:del>/<w:delText>
    and the NEW value in <w:ins>/<w:t>. Accepting changes (the parser's
    accept_revisions) must yield the NEW prices.
    Returns (docx_bytes, {service: (old, new)}).
    """
    names = [
        "Электрофорез", "Магнитотерапия", "Ультразвуковая терапия",
        "Лазеротерапия", "Лечебный массаж спины", "Ингаляция (небулайзер)",
        "УВЧ-терапия",
    ]
    rows = list(_rows_for(names))

    # Pick two rows to revise. Store OLD (in-doc) and NEW (final after accept).
    revisions: dict[str, tuple[str, str]] = {}
    # Massage: massage old 5 000 -> new 8 000 (price increase tracked change).
    revisions["Лечебный массаж спины"] = (_kzt(5000), _kzt(8000))
    # Laser: old 3 000 -> new 4 500.
    revisions["Лазеротерапия"] = (_kzt(3000), _kzt(4500))

    raw = _docx_header_and_table(clinic, d, rows)
    patched = _inject_tracked_changes(raw, revisions)
    # Report the FINAL (accepted) prices.
    final = {k: v[1] for k, v in revisions.items()}
    return patched, final


def _inject_tracked_changes(
    docx_bytes: bytes, revisions: dict[str, tuple[str, str]]
) -> bytes:
    """Rewrite price cells as <w:del>(old) + <w:ins>(new) revisions.

    For each (service, (old_str, new_str)): locate the table row whose first
    cell text == service, then in its 2nd cell (resident price) replace the run
    text so the cell contains a deleted old value and an inserted new value.
    """
    import zipfile as _zip

    with _zip.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        names = zin.namelist()
        contents = {n: zin.read(n) for n in names}

    root = etree.fromstring(contents["word/document.xml"])

    def w(tag: str) -> str:
        return f"{{{_W_NS}}}{tag}"

    nsmap = {"w": _W_NS}
    author = "Прайс-отдел"
    rev_date = "2025-06-15T10:00:00Z"
    rid = 1000

    def make_run(text: str) -> etree._Element:
        r = etree.SubElement(etree.Element("tmp"), w("r"))
        t = etree.SubElement(r, w("t"))
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        return r

    def make_del_run(text: str) -> etree._Element:
        r = etree.Element(w("r"))
        dt = etree.SubElement(r, w("delText"))
        dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        dt.text = text
        return r

    # Iterate table rows; match first-cell text -> patch second cell.
    patched = 0
    for tr in root.findall(f".//{w('tr')}", nsmap):
        tcs = tr.findall(w("tc"))
        if len(tcs) < 2:
            continue
        first_text = "".join(tcs[0].itertext()).strip()
        if first_text not in revisions:
            continue
        old_str, new_str = revisions[first_text]
        price_tc = tcs[1]

        # Find the paragraph holding the price text inside the price cell.
        para = price_tc.find(w("p"))
        if para is None:
            continue

        # Remove existing runs (the old visible text) from the paragraph but
        # keep paragraph properties (<w:pPr>).
        for r in para.findall(w("r")):
            para.remove(r)

        # 1) <w:del> wrapping the OLD value (deletion).
        nonlocal_rid = rid + patched * 2
        del_el = etree.SubElement(para, w("del"))
        del_el.set(w("id"), str(nonlocal_rid))
        del_el.set(w("author"), author)
        del_el.set(w("date"), rev_date)
        del_el.append(make_del_run(old_str))

        # 2) <w:ins> wrapping the NEW value (insertion).
        ins_el = etree.SubElement(para, w("ins"))
        ins_el.set(w("id"), str(nonlocal_rid + 1))
        ins_el.set(w("author"), author)
        ins_el.set(w("date"), rev_date)
        ins_el.append(make_run(new_str))

        patched += 1

    if patched != len(revisions):
        raise RuntimeError(
            f"tracked-changes injection patched {patched}/{len(revisions)} rows"
        )

    contents["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    out = io.BytesIO()
    with _zip.ZipFile(out, "w", _zip.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, contents[n])
    return out.getvalue()


# --------------------------------------------------------------------------- #
# Orchestration.                                                              #
# --------------------------------------------------------------------------- #
def generate_entries() -> tuple[list[tuple[str, bytes]], dict]:
    """Build all archive entries. Returns (entries, report_meta)."""
    sunkar, dostar, shipager, olimp, aktis = CLINICS
    entries: list[tuple[str, bytes]] = []
    report: dict = {"files": {}, "ocr": None, "tracked_final": None}

    # ---- TEXT PDF #1: Sunkar lab/diag, current. -------------------------- #
    d1 = date(2025, 1, 15)
    pdf1_names = [
        "Общий анализ крови", "Общий анализ мочи", "Биохимический анализ крови",
        "Глюкоза крови", "Гликированный гемоглобин", "Липидный профиль",
        "Коагулограмма", "Тиреотропный гормон", "Свободный тироксин",
        "Витамин D (25-OH)", "С-реактивный белок", "Ферритин",
        "УЗИ органов брюшной полости", "УЗИ щитовидной железы",
        "Электрокардиография", "Эхокардиография",
        "Консультация терапевта", "Консультация кардиолога",
    ]
    pdf1_rows = list(_rows_for(pdf1_names))
    # Edge cases:
    pdf1_rows.append(("Креатинин", _kzt(1600), _kzt(1400)))  # nonres < res
    pdf1_rows.append(("Мочевина", "договорная", ""))          # договорная + empty
    # Off-catalog service (no catalog concept) -> UNMATCHED queue.
    pdf1_rows.append(("Криодеструкция папиллом азотом", _kzt(7800), _kzt(9500)))
    fn = _safe_filename(sunkar, "pdf", d1)
    entries.append((fn, build_text_pdf(sunkar, d1, pdf1_rows)))
    report["files"][fn] = (
        "TEXT PDF (reportlab, Cyrillic) | header+БИН+date; edge: nonres<res, "
        "договорная, empty price, off-catalog service (unmatched)"
    )

    # ---- PRICE HISTORY: Sunkar same services, two dates ------------------ #
    # 2024-01 (old) then 2025-06 (new) with one >50% jump on Витамин D.
    d_old = date(2024, 1, 10)
    d_new = date(2025, 6, 20)
    hist_names = [
        "Общий анализ крови", "Глюкоза крови", "Витамин D (25-OH)",
        "Тиреотропный гормон", "УЗИ органов брюшной полости",
    ]
    old_rows = list(_rows_for(hist_names))
    # New version: bump Витамин D from 9000 -> 19000 (>50% jump -> PRICE_ANOMALY)
    new_rows = []
    for nm in hist_names:
        s = _by_name(nm)
        if nm == "Витамин D (25-OH)":
            res = 19000
        else:
            res = int(round(s.base * 1.10))  # modest 10% inflation, no anomaly
        new_rows.append((nm, _kzt(res), _kzt(int(round(res * 1.45)))))

    fo = _safe_filename(sunkar, "pdf", d_old)
    entries.append((fo, build_text_pdf(
        sunkar, d_old, old_rows, title_note="Архивный прайс (предыдущая версия)."
    )))
    report["files"][fo] = (
        "TEXT PDF (reportlab) | HISTORY old version (2024-01) for versioning"
    )

    fnw = _safe_filename(sunkar, "pdf", d_new)
    entries.append((fnw, build_text_pdf(
        sunkar, d_new, new_rows,
        title_note="Обновлённый прайс. Витамин D пересмотрен.",
    )))
    report["files"][fnw] = (
        "TEXT PDF (reportlab) | HISTORY new version (2025-06); Витамин D "
        "9000->19000 (>50% jump -> PRICE_ANOMALY)"
    )

    # ---- SCANNED PDF (image-only OCR): Shipager ------------------------- #
    ds = date(2025, 3, 5)
    scan_names = [
        "Общий анализ крови", "Общий анализ мочи", "Глюкоза крови",
        "УЗИ органов брюшной полости", "УЗИ щитовидной железы",
        "Рентгенография органов грудной клетки", "Электрокардиография",
        "Консультация терапевта", "Консультация невролога",
    ]
    scan_rows = list(_rows_for(scan_names))
    scan_img = render_pricelist_image(shipager, ds, scan_rows)
    scan_pdf = image_to_scanned_pdf(scan_img)

    # Confirm image-only: no selectable text (mirrors detect._classify_pdf).
    with fitz.open(stream=scan_pdf, filetype="pdf") as _d:
        sel = "".join(p.get_text() for p in _d).strip()
    if len(sel) > 5:
        raise RuntimeError(f"Scanned PDF unexpectedly has selectable text: {sel[:80]!r}")

    # OCR the FINAL rasterized PDF page exactly the way pdf_scan.py would, so the
    # verification reflects what the parser will actually see (post-compression).
    ocr_ok, ocr_text = verify_scanned_ocr(
        _rasterize_pdf_page(scan_pdf), expect_tokens=["анализ", "крови", "2"]
    )
    if not ocr_ok:
        raise RuntimeError(
            "Scanned PDF failed OCR token check. OCR sample:\n" + ocr_text[:600]
        )
    fn = _safe_filename(shipager, "pdf", ds)
    fn = fn.replace("прайс", "скан_прайс")
    entries.append((fn, scan_pdf))
    report["files"][fn] = (
        "SCANNED PDF (image-only, fitz.insert_image) | OCR target, NO "
        "selectable text"
    )
    report["ocr"] = {
        "selectable_text_len": len(sel),
        "sample": "\n".join(
            ln for ln in ocr_text.splitlines() if ln.strip()
        )[:500],
    }

    # ---- XLSX #1: Olimp, separate resident/non-resident, header not row 1 - #
    dx = date(2025, 2, 1)
    fn = _safe_filename(olimp, "xlsx", dx)
    entries.append((fn, build_xlsx_two_price(olimp, dx)))
    report["files"][fn] = (
        "XLSX (multi-sheet, header on row 4) | резидент/нерезидент cols; edge: "
        "nonres<res, договорная, empty price, off-catalog (unmatched + "
        "needs_review); synonyms on "
        "Диагностика sheet"
    )

    # ---- XLSX #2: Aktis, includes USD sheet ------------------------------ #
    dx2 = date(2025, 4, 12)
    fn = _safe_filename(aktis, "xlsx", dx2)
    entries.append((fn, build_xlsx_usd(aktis, dx2)))
    report["files"][fn] = (
        "XLSX (multi-sheet) | KZT sheet + USD sheet ('Цена, USD') for currency "
        "conversion; off-catalog USD row (unmatched)"
    )

    # ---- DOCX #1: Dostar, plain ----------------------------------------- #
    dd = date(2025, 1, 28)
    fn = _safe_filename(dostar, "docx", dd)
    entries.append((fn, build_docx_plain(dostar, dd)))
    report["files"][fn] = (
        "DOCX (python-docx) | consultations + ФД; edge: nonres<res"
    )

    # ---- DOCX #2: Dostar, TRACKED CHANGES -------------------------------- #
    dd2 = date(2025, 6, 15)
    docx_tracked, final_prices = build_docx_tracked(dostar, dd2)
    fn = _safe_filename(dostar, "docx", dd2)
    fn = fn.replace("прайс", "прайс_правки")
    entries.append((fn, docx_tracked))
    report["files"][fn] = (
        "DOCX with TRACKED CHANGES (<w:ins>/<w:del>) | accepting yields NEW "
        f"prices: {final_prices}"
    )
    report["tracked_final"] = final_prices

    return entries, report


def write_archive(entries: list[tuple[str, bytes]], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in entries:
            z.writestr(name, data)


# --------------------------------------------------------------------------- #
# Sanity readback.                                                            #
# --------------------------------------------------------------------------- #
def sanity_readback(path: Path) -> list[str]:
    """Open a few entries back with their libs to confirm they're well-formed."""
    notes: list[str] = []
    with zipfile.ZipFile(path, "r") as z:
        names = z.namelist()

        # An XLSX: open with openpyxl.
        xlsx_name = next(n for n in names if n.endswith(".xlsx"))
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(z.read(xlsx_name)), read_only=True)
        notes.append(
            f"xlsx OK: {xlsx_name} sheets={wb.sheetnames}"
        )
        wb.close()

        # A DOCX: open with python-docx (the tracked one) + accept revisions.
        docx_name = next(n for n in names if n.endswith(".docx") and "правки" in n)
        tmp = OUT_DIR / "_tmp_readback.docx"
        tmp.write_bytes(z.read(docx_name))
        try:
            from app.parsers.docx_parser import accept_revisions

            cleaned = accept_revisions(str(tmp))
            from docx import Document as _Doc

            cdoc = _Doc(io.BytesIO(cleaned))
            cell_texts = []
            for t in cdoc.tables:
                for row in t.rows:
                    cell_texts.append([c.text.strip() for c in row.cells])
            notes.append(
                f"docx OK (revisions accepted): {docx_name}; rows={len(cell_texts)}"
            )
            # Verify accepted prices are the NEW ones.
            flat = {r[0]: r[1] for r in cell_texts if len(r) >= 2}
            notes.append(
                "  accepted massage price -> "
                f"{flat.get('Лечебный массаж спины')!r} (expect 8 000)"
            )
            notes.append(
                "  accepted laser price   -> "
                f"{flat.get('Лазеротерапия')!r} (expect 4 500)"
            )
        finally:
            tmp.unlink(missing_ok=True)

        # A text PDF: open with fitz, confirm selectable text present.
        pdf_text_name = next(
            n for n in names if n.endswith(".pdf") and "скан" not in n
        )
        with fitz.open(stream=z.read(pdf_text_name), filetype="pdf") as d:
            txt = "".join(p.get_text() for p in d)
        notes.append(
            f"pdf(text) OK: {pdf_text_name}; selectable_chars={len(txt.split())} words"
        )

        # The scanned PDF: confirm image-only.
        scan_name = next(n for n in names if "скан" in n)
        with fitz.open(stream=z.read(scan_name), filetype="pdf") as d:
            txt = "".join(p.get_text() for p in d).strip()
        notes.append(
            f"pdf(scan) OK: {scan_name}; selectable_text_len={len(txt)} (expect ~0)"
        )
    return notes


def main() -> None:
    entries, report = generate_entries()
    write_archive(entries, ARCHIVE_PATH)

    print(f"Wrote {ARCHIVE_PATH}")
    print(f"\n=== ARCHIVE MANIFEST ({len(entries)} entries) ===")
    with zipfile.ZipFile(ARCHIVE_PATH, "r") as z:
        for info in z.infolist():
            note = report["files"].get(info.filename, "")
            print(f"  {info.file_size:>8} B  {info.filename}")
            if note:
                print(f"             -> {note}")

    print("\n=== SCANNED PDF / OCR ===")
    print(f"  selectable text length: {report['ocr']['selectable_text_len']} (image-only)")
    print("  OCR sample (first lines):")
    for ln in report["ocr"]["sample"].splitlines()[:10]:
        print(f"    | {ln}")

    print("\n=== TRACKED CHANGES (accepted/final prices) ===")
    print(f"  {report['tracked_final']}")

    print("\n=== SANITY READBACK ===")
    for n in sanity_readback(ARCHIVE_PATH):
        print(f"  {n}")


if __name__ == "__main__":
    main()
