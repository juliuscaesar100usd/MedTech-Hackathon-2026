"""DOCX parser with tracked-changes acceptance (§4.2).

Word price lists may contain unaccepted revisions. We MUST work with the final
(accepted) text: every <w:del> (deleted content) is removed entirely, and every
<w:ins> (inserted content) is unwrapped so its runs survive. The cleaned XML is
then loaded into python-docx for table + paragraph extraction.
"""
from __future__ import annotations

import re
import zipfile
from io import BytesIO

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from ..enums import FileFormat
from .base import BaseParser, ParsedDocument, register_parser
from .pdf_text import extract_header_hints
from .table_extract import (
    SectionHierarchy,
    _section_depth,
    _section_label,
    rows_from_table,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DOC_XML = "word/document.xml"

# Word built-in heading styles: "Heading 1".."Heading 9" / localized "Заголовок N".
_HEADING_STYLE_RE = re.compile(r"(?:heading|заголовок)\s*(\d+)", re.I)
# python-docx left_indent is in EMU; ~457200 EMU = 0.5 inch ≈ one indent level.
_EMU_PER_LEVEL = 457200


def accept_revisions(path: str) -> bytes:
    """Return the .docx bytes with all tracked changes accepted.

    - <w:del>/<w:delText> removed (deletions vanish),
    - <w:ins> unwrapped (insertions kept),
    - move/format-change markers (<w:moveFrom>, <w:rPrChange>...) handled so the
      final text is clean.
    """
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        contents = {n: zin.read(n) for n in names}

    xml = contents.get(_DOC_XML)
    if xml is None:
        return _repack(contents, names)  # not a Word doc body; pass through

    root = etree.fromstring(xml)

    def w(tag: str) -> str:
        return f"{{{_W_NS}}}{tag}"

    # 1) Drop deleted content + move-source content entirely.
    for tag in ("del", "moveFrom"):
        for el in root.findall(f".//{w(tag)}"):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
    # delText/delInstrText that somehow survived outside a <w:del>
    for tag in ("delText", "delInstrText"):
        for el in root.findall(f".//{w(tag)}"):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    # 2) Unwrap inserted/moved-in content: replace the wrapper with its children.
    for tag in ("ins", "moveTo"):
        for el in root.findall(f".//{w(tag)}"):
            parent = el.getparent()
            if parent is None:
                continue
            idx = parent.index(el)
            for child in reversed(list(el)):
                parent.insert(idx, child)
            parent.remove(el)

    # 3) Drop format-/property-change tracking markers (keep current props).
    for tag in ("rPrChange", "pPrChange", "tblPrChange", "trPrChange",
                "tcPrChange", "sectPrChange", "numberingChange"):
        for el in root.findall(f".//{w(tag)}"):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    contents[_DOC_XML] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    return _repack(contents, names)


def _repack(contents: dict[str, bytes], order: list[str]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in order:
            zout.writestr(name, contents[name])
    return buf.getvalue()


def _table_to_2d(table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _iter_block_items(document):
    """Yield paragraphs and tables in DOCUMENT ORDER.

    python-docx exposes ``document.paragraphs`` and ``document.tables`` as two
    separate flat lists, losing their interleaving. Price lists routinely put a
    section heading in a paragraph just above the table it labels, so we walk the
    body XML directly to preserve order.
    """
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _style_depth(paragraph) -> int | None:
    """Nesting depth (1-based) if `paragraph` is a styled/structural heading.

    Heading 1..9 -> 1..9; Title -> 1; a numbered/bulleted list level -> ilvl+1;
    an explicit left indent on a short line -> indent buckets. Ordinary body
    paragraphs return None.
    """
    try:
        style_name = paragraph.style.name or ""
    except Exception:
        style_name = ""
    m = _HEADING_STYLE_RE.search(style_name)
    if m:
        return int(m.group(1))
    if style_name.strip().lower() in ("title", "название", "заголовок"):
        return 1
    # Numbered / bulleted list nesting level (w:pPr/w:numPr/w:ilvl).
    try:
        pPr = paragraph._p.pPr
        if pPr is not None and pPr.numPr is not None and pPr.numPr.ilvl is not None:
            val = pPr.numPr.ilvl.val
            if val is not None:
                return int(val) + 1
    except Exception:
        pass
    # Explicit left indent — only trust it on short, heading-like lines so a
    # normal indented sentence isn't mistaken for a section header.
    try:
        li = paragraph.paragraph_format.left_indent
        text = (paragraph.text or "").strip()
        if li is not None and int(li) > 0 and len(text) <= 60:
            return 1 + int(li) // _EMU_PER_LEVEL
    except Exception:
        pass
    return None


def _paragraph_section(paragraph) -> tuple[str, int] | None:
    """Return (label, depth) if a paragraph is a section heading, else None.

    Style/list/indent give the depth directly (1-based). For an UNstyled line we
    only treat it as a section when it reads like one (reusing the table heading
    heuristics), mapping the coarse keyword class onto the same 1-based scale.
    """
    text = (paragraph.text or "").strip()
    if len(text) < 2 or not any(ch.isalpha() for ch in text):
        return None
    depth = _style_depth(paragraph)
    if depth is not None:
        return text.strip(" .:-—|\t"), depth
    label = _section_label([text])
    if label is None:
        return None
    # _section_depth returns 0..2 here (no geometry); shift onto the 1-based
    # heading scale so it composes with style-derived depths.
    return label, _section_depth(label) + 1


@register_parser
class DocxParser(BaseParser):
    formats = (FileFormat.docx,)

    def parse(self, file_path: str, file_name: str | None = None) -> ParsedDocument:
        doc = ParsedDocument(file_format=FileFormat.docx)
        try:
            cleaned = accept_revisions(file_path)
            document = Document(BytesIO(cleaned))
        except Exception as exc:
            doc.add_warning(f"DOCX load failed: {exc}")
            return doc

        # Walk paragraphs + tables in document order so heading paragraphs above
        # a table flow into that table's rows as a nested section_path. Paragraph
        # headings build a hierarchy; each table inherits its current path as a
        # base_path, with any IN-table section labels nesting beneath it.
        hierarchy = SectionHierarchy()
        tno = 0
        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                sec = _paragraph_section(block)
                if sec is not None:
                    hierarchy.push(sec[0], sec[1])
            elif isinstance(block, Table):
                tno += 1
                doc.rows.extend(
                    rows_from_table(
                        _table_to_2d(block),
                        source_ref_prefix=f"table={tno}",
                        base_path=hierarchy.path(),
                    )
                )

        paras = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        doc.raw_text = "\n".join(paras)
        if not doc.rows:
            doc.add_warning("No tables extracted from DOCX.")
        extract_header_hints(doc.raw_text, doc)
        return doc
