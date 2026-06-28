"""Tests for hierarchical section parsing (section_path on ParsedRow).

Covers, per the hierarchy contract:
  1. the SectionHierarchy stack push/pop-by-depth logic,
  2. the generic keyword-class depth fallback (department -> group -> leaf),
  3. rows_from_table wiring (level_hints + base_path -> section_path),
  4. real end-to-end parses that yield multi-level nesting:
       * a DOCX with Heading 1/2/3 paragraphs above tables (REQUIRED), and
       * an XLSX whose cell-alignment indent drives the nesting.
"""
from __future__ import annotations

import openpyxl
import pytest
from openpyxl.styles import Alignment

from app.enums import Currency, FileFormat
from app.parsers import parse_file
from app.parsers.base import get_parser
from app.parsers.table_extract import (
    SectionHierarchy,
    _generic_section_depth,
    _section_depth,
    rows_from_table,
)


# --------------------------------------------------------------------------- #
# (1) SectionHierarchy stack logic.                                            #
# --------------------------------------------------------------------------- #
def test_hierarchy_deeper_nests_same_level_replaces():
    h = SectionHierarchy()
    h.push("Лаборатория", 0)
    assert h.path() == ["Лаборатория"]
    assert h.innermost == "Лаборатория"

    h.push("Анализ крови", 1)          # strictly deeper -> nests
    assert h.path() == ["Лаборатория", "Анализ крови"]

    h.push("Гормоны", 2)               # deeper still -> nests
    assert h.path() == ["Лаборатория", "Анализ крови", "Гормоны"]
    assert h.innermost == "Гормоны"
    assert len(h) == 3

    h.push("Биохимия", 2)              # same level -> replaces the sibling
    assert h.path() == ["Лаборатория", "Анализ крови", "Биохимия"]

    h.push("Диагностика", 0)           # shallower -> pops everything below it
    assert h.path() == ["Диагностика"]
    assert len(h) == 1


def test_hierarchy_pops_to_intermediate_level():
    h = SectionHierarchy()
    for label, depth in (("A", 0), ("B", 1), ("C", 2)):
        h.push(label, depth)
    assert h.path() == ["A", "B", "C"]
    # A header at depth 1 pops C(2) and B(1), then nests under A(0).
    h.push("B2", 1)
    assert h.path() == ["A", "B2"]


def test_hierarchy_reset():
    h = SectionHierarchy()
    h.push("X", 0)
    h.reset()
    assert h.path() == []
    assert h.innermost is None
    assert len(h) == 0


# --------------------------------------------------------------------------- #
# (2) Generic keyword-class depth fallback.                                    #
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    "label,expected",
    [
        ("Лаборатория", 0),                  # department keyword
        ("Диагностика", 0),                  # department keyword
        ("ЛАБОРАТОРНЫЕ ИССЛЕДОВАНИЯ", 0),     # ALL-CAPS banner -> top level
        ("Анализ крови", 1),                 # group word
        ("Биохимическая панель", 1),         # group word ("панель")
        ("Прочее:", 1),                      # ':'-suffixed header
        ("Гормоны", 2),                      # no keyword -> leaf-ish sub-group
        ("Витамин D", 2),
    ],
)
def test_generic_section_depth(label, expected):
    assert _generic_section_depth(label) == expected


def test_section_depth_geometry_wins_over_keyword():
    # A geometry hint always overrides the keyword guess.
    assert _section_depth("Гормоны", geo=5) == 5
    assert _section_depth("Лаборатория", geo=3) == 3
    # No hint -> keyword class.
    assert _section_depth("Гормоны") == 2
    assert _section_depth("Лаборатория") == 0


# --------------------------------------------------------------------------- #
# (3) rows_from_table: level_hints + base_path wiring.                         #
# --------------------------------------------------------------------------- #
def test_rows_from_table_generic_two_level_nesting():
    # Only headers that _section_label recognizes (ALL-CAPS banners or keyword
    # rows) can form a level. Generic keyword class then gives department=0 and
    # group=1 -> a two-level nesting, with same-level siblings replacing.
    table = [
        ["Наименование", "Цена"],
        ["ЛАБОРАТОРНЫЕ ИССЛЕДОВАНИЯ"],    # ALL-CAPS banner -> 0
        ["Анализ крови"],                 # group -> 1
        ["Гемоглобин", "500"],
        ["Биохимические исследования"],   # group -> 1, replaces the sibling
        ["Глюкоза", "800"],
        ["ДИАГНОСТИКА"],                  # ALL-CAPS banner -> 0, pops everything
        ["УЗИ брюшной полости", "8 000"],
    ]
    rows = rows_from_table(table)
    by_name = {r.service_name_raw: r for r in rows}

    assert by_name["Гемоглобин"].section_path == [
        "ЛАБОРАТОРНЫЕ ИССЛЕДОВАНИЯ", "Анализ крови",
    ]
    assert by_name["Гемоглобин"].extra["section"] == "Анализ крови"  # innermost
    assert by_name["Глюкоза"].section_path == [
        "ЛАБОРАТОРНЫЕ ИССЛЕДОВАНИЯ", "Биохимические исследования",
    ]
    assert by_name["УЗИ брюшной полости"].section_path == ["ДИАГНОСТИКА"]


def test_rows_from_table_level_hints_drive_nesting():
    # Both "Терапия" and "Хирургия" are department keywords (generic depth 0),
    # so WITHOUT a hint the second would replace the first. A geometry hint of 2
    # on "Хирургия" must make it nest instead -> proves level_hints is used.
    table = [
        ["Наименование", "Цена"],
        ["Терапия"],
        ["Хирургия"],
        ["Прием хирурга", "6 000"],
    ]
    hints = [None, None, 2, None]  # aligned with table row indices
    rows = rows_from_table(table, level_hints=hints)
    assert rows[0].section_path == ["Терапия", "Хирургия"]
    assert rows[0].extra["section"] == "Хирургия"

    # Sanity: without the hint they collapse to a single replaced sibling.
    rows_flat = rows_from_table(table)
    assert rows_flat[0].section_path == ["Хирургия"]


def test_rows_from_table_base_path_prefix():
    # base_path (e.g. DOCX paragraph headings) is prepended; in-table sections
    # nest beneath it.
    rows = rows_from_table(
        [["Наименование", "Цена"], ["ТТГ", "3 500"]],
        base_path=["Лаборатория", "Гормоны"],
    )
    assert rows[0].section_path == ["Лаборатория", "Гормоны"]
    assert rows[0].extra["section"] == "Гормоны"

    rows2 = rows_from_table(
        [["Наименование", "Цена"], ["Анализ крови"], ["Глюкоза", "800"]],
        base_path=["Лаборатория"],
    )
    assert rows2[0].section_path == ["Лаборатория", "Анализ крови"]


def test_rows_from_table_no_section_is_empty_path():
    rows = rows_from_table([["Наименование", "Цена"], ["ЭКГ", "3 500"]])
    assert rows[0].section_path == []
    assert "section" not in rows[0].extra


# --------------------------------------------------------------------------- #
# (4a) End-to-end XLSX: cell-alignment indent drives the nesting.             #
# --------------------------------------------------------------------------- #
def _make_nested_xlsx(path: str) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Прайс"
    ws.append(["Наименование", "Цена, тг"])      # row 1: column header
    ws.append(["Терапия", None])                 # row 2: dept, no indent
    ws.append(["Хирургия", None])                # row 3: indented -> nests
    ws["A3"].alignment = Alignment(indent=2)
    ws.append(["Прием хирурга", "6 000"])        # row 4: data
    wb.save(path)


def test_xlsx_indent_builds_section_path(tmp_path):
    path = str(tmp_path / "nested.xlsx")
    _make_nested_xlsx(path)
    doc = parse_file(path)
    assert doc.file_format == FileFormat.xlsx
    by_name = {r.service_name_raw: r for r in doc.rows}
    assert "Прием хирурга" in by_name
    row = by_name["Прием хирурга"]
    # "Терапия" and "Хирургия" are both department keywords (generic depth 0);
    # only the alignment indent on "Хирургия" makes it nest rather than replace.
    assert row.section_path == ["Терапия", "Хирургия"]
    assert row.extra["section"] == "Хирургия"
    assert row.price_resident == 6000.0
    assert row.currency == Currency.KZT


# --------------------------------------------------------------------------- #
# (4b) End-to-end DOCX: Heading 1/2/3 paragraphs above tables (REQUIRED).      #
# --------------------------------------------------------------------------- #
def _make_nested_docx(path: str) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Лаборатория", level=1)
    document.add_heading("Анализ крови", level=2)
    document.add_heading("Гормоны", level=3)
    t1 = document.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "Наименование"
    t1.cell(0, 1).text = "Цена, тг"
    t1.cell(1, 0).text = "ТТГ"
    t1.cell(1, 1).text = "3 500"

    # New department -> must pop the lab chain entirely.
    document.add_heading("Диагностика", level=1)
    document.add_heading("УЗИ", level=2)
    t2 = document.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "Наименование"
    t2.cell(0, 1).text = "Цена, тг"
    t2.cell(1, 0).text = "УЗИ органов малого таза"
    t2.cell(1, 1).text = "8 000"
    document.save(path)


def test_docx_heading_paragraphs_build_nested_path(tmp_path):
    path = str(tmp_path / "nested.docx")
    _make_nested_docx(path)
    parser = get_parser(FileFormat.docx)
    doc = parser.parse(path)
    assert doc.file_format == FileFormat.docx
    by_name = {r.service_name_raw: r for r in doc.rows}

    assert "ТТГ" in by_name, by_name.keys()
    ttg = by_name["ТТГ"]
    assert ttg.section_path == ["Лаборатория", "Анализ крови", "Гормоны"]
    assert ttg.extra["section"] == "Гормоны"
    assert ttg.price_resident == 3500.0

    uzi = by_name["УЗИ органов малого таза"]
    assert uzi.section_path == ["Диагностика", "УЗИ"]
    assert uzi.extra["section"] == "УЗИ"
