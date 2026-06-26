"""Tests for the MedArchive PARSERS module.

Fixtures are built programmatically in a tmp dir so the suite is hermetic:
  1. text PDF (reportlab) -> PdfTextParser rows + hints,
  2. xlsx (openpyxl) with title rows before the header -> header detection,
  3. docx (python-docx) with an injected tracked change -> accepted value wins,
  4. scanned PDF (PIL image embedded via fitz) -> OCR extracts the known row.
"""
from __future__ import annotations

import os
import zipfile
from io import BytesIO

import pytest

from app.enums import Currency, FileFormat
from app.parsers import detect_format, parse_file
from app.parsers.base import get_parser
from app.parsers.table_extract import find_header_row, parse_price, rows_from_table


# Cyrillic-capable TrueType font for building the sample PDFs/scans. The Linux
# CI image ships DejaVuSans at the first path; macOS/Windows dev machines don't,
# so we fall back to system fonts that also cover Cyrillic. Resolved once.
_CYRILLIC_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",          # Debian / Ubuntu (CI)
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",                   # Fedora / Arch
    "/Library/Fonts/Arial Unicode.ttf",                         # macOS
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",     # macOS
    "/System/Library/Fonts/Supplemental/Arial.ttf",            # macOS
    "/System/Library/Fonts/Supplemental/Verdana.ttf",          # macOS
    "C:\\Windows\\Fonts\\arial.ttf",                            # Windows
)


def _cyrillic_font_path() -> str:
    """First Cyrillic-capable TTF that exists, or skip on a fontless system."""
    for p in _CYRILLIC_FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    pytest.skip("no Cyrillic-capable TrueType font found on this system")


# --------------------------------------------------------------------------- #
# Unit tests: price parsing + header detection + table mapping.               #
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize(
    "text,expected_price,expected_cur",
    [
        ("12 000", 12000.0, Currency.KZT),
        ("12000,00", 12000.0, Currency.KZT),
        ("12 000 ₸", 12000.0, Currency.KZT),
        ("15 000 тг", 15000.0, Currency.KZT),
        ("$120", 120.0, Currency.USD),
        ("120 USD", 120.0, Currency.USD),
        ("1 500 руб", 1500.0, Currency.RUB),
        ("—", None, Currency.KZT),
        ("-", None, Currency.KZT),
        ("договорная", None, Currency.KZT),
        ("", None, Currency.KZT),
    ],
)
def test_parse_price(text, expected_price, expected_cur):
    price, cur = parse_price(text)
    assert price == expected_price
    assert cur == expected_cur


def test_find_header_row_skips_titles():
    table = [
        ["ТОО Медцентр Альфа"],
        ["г. Алматы"],
        ["Наименование услуги", "Код", "Цена, тг"],
        ["Общий анализ крови", "A01", "2 500"],
    ]
    assert find_header_row(table) == 2


def test_rows_from_table_maps_columns():
    table = [
        ["Наименование", "Код", "Цена (резидент)", "Цена (нерезидент)"],
        ["УЗИ брюшной полости", "U10", "8 000 тг", "12 000 тг"],
        ["Договорная позиция", "X", "договорная", "-"],
    ]
    rows = rows_from_table(table, source_ref_prefix="sheet=Прайс")
    assert len(rows) == 2
    r = rows[0]
    assert r.service_name_raw == "УЗИ брюшной полости"
    assert r.service_code_source == "U10"
    assert r.price_resident == 8000.0
    assert r.price_nonresident == 12000.0
    assert r.currency == Currency.KZT
    assert r.source_ref.startswith("sheet=Прайс;row=")
    # second row keeps the name but has no price
    assert rows[1].price_resident is None
    assert rows[1].price_nonresident is None


# --------------------------------------------------------------------------- #
# Fixture builders.                                                            #
# --------------------------------------------------------------------------- #
def _make_text_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    # Helvetica has no Cyrillic glyphs; register a TTF that does.
    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", _cyrillic_font_path()))

    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    y = h - 30 * mm
    c.setFont("DejaVu", 14)
    c.drawString(25 * mm, y, "ТОО Медицинский центр Альфа")
    y -= 8 * mm
    c.setFont("DejaVu", 10)
    c.drawString(25 * mm, y, "г. Алматы, ул. Абая 10")
    y -= 6 * mm
    c.drawString(25 * mm, y, "БИН: 123456789012  тел: +7 701 234 56 78")
    y -= 6 * mm
    c.drawString(25 * mm, y, "Прайс-лист на 01.03.2026  email: info@alpha-clinic.kz")
    y -= 12 * mm

    # A simple bordered table.
    from reportlab.platypus import Table

    data = [
        ["Наименование услуги", "Код", "Цена, тг"],
        ["Общий анализ крови", "A01", "2 500"],
        ["УЗИ брюшной полости", "U10", "8 000"],
        ["Консультация терапевта", "C05", "6 000"],
    ]
    t = Table(data, colWidths=[70 * mm, 25 * mm, 30 * mm])
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    t.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "DejaVu"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    tw, th = t.wrap(0, 0)
    t.drawOn(c, 25 * mm, y - th)
    c.showPage()
    c.save()


def _make_xlsx(path: str) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Прайс"
    ws.append(["ТОО Медцентр Бета"])               # title row 1
    ws.append(["г. Астана, БИН: 987654321098"])    # title row 2
    ws.append([])                                   # blank
    ws.append(["Наименование", "Код", "Стоимость, тг"])  # header row 3 (idx 3)
    ws.append(["ЭКГ", "E01", "3 500"])
    ws.append(["Рентген грудной клетки", "R02", "5 000"])
    wb.save(path)


def _make_docx_with_tracked_change(path: str) -> None:
    """Build a docx, then inject a tracked change into word/document.xml:
    the price cell shows old '5 000' (deleted) and new '7 500' (inserted).
    """
    from docx import Document

    document = Document()
    document.add_heading("ТОО Клиника Гамма", level=1)
    document.add_paragraph("г. Шымкент")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Наименование"
    table.cell(0, 1).text = "Цена, тг"
    table.cell(1, 0).text = "Прием врача"
    table.cell(1, 1).text = "PRICE_PLACEHOLDER"
    document.save(path)

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    # A run with a deleted old price and an inserted new price.
    revised = (
        f'<w:del w:id="1" w:author="a" xmlns:w="{W}">'
        f'<w:r><w:delText>5 000</w:delText></w:r></w:del>'
        f'<w:ins w:id="2" w:author="a" xmlns:w="{W}">'
        f'<w:r><w:t>7 500</w:t></w:r></w:ins>'
    )
    placeholder_run = '<w:r><w:t>PRICE_PLACEHOLDER</w:t></w:r>'

    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        contents = {n: zin.read(n) for n in names}
    xml = contents["word/document.xml"].decode("utf-8")
    assert placeholder_run in xml, "placeholder run not found in document.xml"
    xml = xml.replace(placeholder_run, revised)
    contents["word/document.xml"] = xml.encode("utf-8")
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, contents[n])
    with open(path, "wb") as f:
        f.write(buf.getvalue())


def _make_scan_pdf(path: str) -> None:
    """Render a price list onto a white PIL image, embed it full-page in a PDF."""
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (1654, 1169), "white")  # ~A4 landscape @ 200dpi
    draw = ImageDraw.Draw(img)
    font_path = _cyrillic_font_path()
    big = ImageFont.truetype(font_path, 48)
    med = ImageFont.truetype(font_path, 40)

    draw.text((80, 60), "Медцентр Дельта", fill="black", font=big)
    draw.text((80, 160), "Прайс-лист", fill="black", font=med)
    rows = [
        ("Наименование", "Цена"),
        ("Анализ крови", "2500"),
        ("Консультация", "6000"),
    ]
    y = 280
    for name, price in rows:
        draw.text((80, y), name, fill="black", font=med)
        draw.text((1100, y), price, fill="black", font=med)
        y += 90

    png = BytesIO()
    img.save(png, format="PNG")
    png.seek(0)

    doc = fitz.open()
    page = doc.new_page(width=842, height=595)  # A4 landscape pts
    page.insert_image(page.rect, stream=png.read())
    doc.save(path)
    doc.close()


# --------------------------------------------------------------------------- #
# Fixtures.                                                                    #
# --------------------------------------------------------------------------- #
@pytest.fixture(scope="module")
def fixtures(tmp_path_factory):
    d = tmp_path_factory.mktemp("medarchive_fixtures")
    paths = {
        "pdf": str(d / "price_text.pdf"),
        "xlsx": str(d / "price.xlsx"),
        "docx": str(d / "price.docx"),
        "scan": str(d / "price_scan.pdf"),
    }
    _make_text_pdf(paths["pdf"])
    _make_xlsx(paths["xlsx"])
    _make_docx_with_tracked_change(paths["docx"])
    _make_scan_pdf(paths["scan"])
    return paths


# --------------------------------------------------------------------------- #
# Detection tests.                                                            #
# --------------------------------------------------------------------------- #
def test_detect_format_text_pdf(fixtures):
    assert detect_format(fixtures["pdf"]) == FileFormat.pdf


def test_detect_format_xlsx(fixtures):
    assert detect_format(fixtures["xlsx"]) == FileFormat.xlsx


def test_detect_format_docx(fixtures):
    assert detect_format(fixtures["docx"]) == FileFormat.docx


def test_detect_format_scan_pdf(fixtures):
    assert detect_format(fixtures["scan"]) == FileFormat.scan_pdf


# --------------------------------------------------------------------------- #
# (1) Text PDF.                                                                #
# --------------------------------------------------------------------------- #
def test_pdf_text_parser(fixtures):
    doc = parse_file(fixtures["pdf"])
    assert doc.file_format == FileFormat.pdf
    by_name = {r.service_name_raw: r for r in doc.rows}
    assert "Общий анализ крови" in by_name
    assert by_name["Общий анализ крови"].price_resident == 2500.0
    assert by_name["УЗИ брюшной полости"].price_resident == 8000.0
    assert by_name["Консультация терапевта"].price_resident == 6000.0
    assert by_name["Общий анализ крови"].service_code_source == "A01"
    # header hints
    assert doc.partner_name_hint and "Альфа" in doc.partner_name_hint
    assert doc.bin_hint == "123456789012"
    assert doc.email_hint == "info@alpha-clinic.kz"
    assert doc.effective_date_hint is not None
    assert doc.effective_date_hint.year == 2026 and doc.effective_date_hint.month == 3


# --------------------------------------------------------------------------- #
# (2) XLSX with title rows before header.                                      #
# --------------------------------------------------------------------------- #
def test_xlsx_parser_header_detection(fixtures):
    doc = parse_file(fixtures["xlsx"])
    assert doc.file_format == FileFormat.xlsx
    by_name = {r.service_name_raw: r for r in doc.rows}
    assert set(by_name) >= {"ЭКГ", "Рентген грудной клетки"}
    assert by_name["ЭКГ"].price_resident == 3500.0
    assert by_name["ЭКГ"].service_code_source == "E01"
    assert by_name["Рентген грудной клетки"].price_resident == 5000.0
    # title rows must NOT have become service rows
    assert "ТОО Медцентр Бета" not in by_name
    assert all(r.source_ref.startswith("sheet=Прайс") for r in doc.rows)


# --------------------------------------------------------------------------- #
# (3) DOCX tracked changes -> accepted value wins.                            #
# --------------------------------------------------------------------------- #
def test_docx_accepts_tracked_changes(fixtures):
    doc = parse_file(fixtures["docx"])
    assert doc.file_format == FileFormat.docx
    by_name = {r.service_name_raw: r for r in doc.rows}
    assert "Прием врача" in by_name
    row = by_name["Прием врача"]
    # The inserted value 7 500 is kept; the deleted 5 000 is gone.
    assert row.price_resident == 7500.0
    assert "5 000" not in doc.raw_text
    assert "5000" not in (str(row.price_resident))


# --------------------------------------------------------------------------- #
# (4) Scanned PDF -> OCR.                                                      #
# --------------------------------------------------------------------------- #
def test_pdf_scan_parser_ocr(fixtures):
    parser = get_parser(FileFormat.scan_pdf)
    doc = parser.parse(fixtures["scan"])
    assert doc.file_format == FileFormat.scan_pdf
    assert doc.used_ocr is True
    assert doc.language
    text = doc.raw_text.lower()
    # Tolerant assertions: OCR noise is expected, check key substrings/prices.
    assert "крови" in text or "анализ" in text
    assert "2500" in text.replace(" ", "") or "2500" in text
    # At least one priced row should be reconstructed.
    prices = [r.price_resident for r in doc.rows if r.price_resident is not None]
    assert any(p in (2500.0, 6000.0) for p in prices)
